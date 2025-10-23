# backend/documentos.py
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Diretório de saída dos documentos
OUTPUT_DIR = os.path.join(os.getcwd(), "docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def gerar_peticao_word(conteudo: str) -> str:
    """
    Gera um arquivo Word (.docx) com o conteúdo jurídico fornecido.
    Retorna o caminho completo do arquivo gerado.
    """

    doc = Document()

    # ====== ESTILO DO CABEÇALHO ======
    titulo = doc.add_paragraph("Babix IA – Advogada Digital de Trânsito")
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = titulo.runs[0]
    run_titulo.bold = True
    run_titulo.font.size = Pt(16)
    run_titulo.font.color.rgb = RGBColor(178, 34, 34)  # vermelho escuro

    subtitulo = doc.add_paragraph("Geração automática de petições e defesas administrativas")
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitulo.runs[0]
    run_sub.italic = True
    run_sub.font.size = Pt(10)
    run_sub.font.color.rgb = RGBColor(90, 90, 90)

    doc.add_paragraph("")  # espaçamento

    # ====== CONTEÚDO PRINCIPAL ======
    partes = conteudo.split("\n")
    for linha in partes:
        p = doc.add_paragraph(linha.strip())
        p.style = "Normal"
        p_format = p.paragraph_format
        p_format.space_after = Pt(6)
        p_format.line_spacing = 1.2
        run = p.runs[0]
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

    # ====== RODAPÉ ======
    doc.add_paragraph("")
    rodape = doc.add_paragraph(
        f"📅 Documento gerado automaticamente em {datetime.now().strftime('%d/%m/%Y %H:%M')}.\n"
        "👩🏻‍⚖️ Sistema: Babix IA • Direito de Trânsito • MG Multas"
    )
    rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rodape.runs[0].italic = True
    rodape.runs[0].font.size = Pt(9)
    rodape.runs[0].font.color.rgb = RGBColor(128, 0, 0)

    # ====== SALVAR DOCUMENTO ======
    nome_arquivo = f"peticao_babix_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    caminho_completo = os.path.join(OUTPUT_DIR, nome_arquivo)
    doc.save(caminho_completo)

    return caminho_completo
