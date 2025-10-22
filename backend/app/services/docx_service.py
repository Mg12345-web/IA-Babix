
from docx import Document
from docx.shared import Pt
from datetime import datetime
import os, uuid
from pydantic import BaseModel
from typing import List

class DocPayload(BaseModel):
    fatos: str
    direito: str
    jurisprudencias: List[str] = []
    pedidos: List[str] = []
    conclusao: str = "Termos em que, pede deferimento."

def build_docx(payload: DocPayload) -> str:
    doc = Document()
    styles = doc.styles["Normal"]
    styles.font.name = "Times New Roman"
    styles.font.size = Pt(12)

    doc.add_heading("DEFESA — Trânsito", level=1)
    doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y')}")

    doc.add_heading("I. DOS FATOS", level=2)
    doc.add_paragraph(payload.fatos)

    doc.add_heading("II. DO DIREITO", level=2)
    doc.add_paragraph(payload.direito)

    if payload.jurisprudencias:
        doc.add_heading("III. JURISPRUDÊNCIAS", level=2)
        for j in payload.jurisprudencias:
            doc.add_paragraph(f"- {j}")

    if payload.pedidos:
        doc.add_heading("IV. DOS PEDIDOS", level=2)
        for p in payload.pedidos:
            doc.add_paragraph(f"- {p}")

    doc.add_heading("V. CONCLUSÃO", level=2)
    doc.add_paragraph(payload.conclusao)

    outdir = "/mnt/data/output"
    os.makedirs(outdir, exist_ok=True)
    path = os.path.join(outdir, f"defesa_{uuid.uuid4().hex}.docx")
    doc.save(path)
    return path
